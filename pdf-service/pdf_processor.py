from pdf2docx import Converter
from docx import Document
from docx.oxml.ns import qn
import os
import zipfile
import re
from uuid import uuid4


class PDFProcessor:

    def process_pdf(self, pdf_path):

        docx_path = pdf_path.replace(
            ".pdf",
            f"_{uuid4().hex}.docx"
        )

        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()

        doc = Document(docx_path)

        # =====================================================
        # TABLE EXTRACTION
        # =====================================================

        extracted_tables = []

        print("\nTABLES FOUND:")
        print(len(doc.tables))

        for table_index, table in enumerate(doc.tables):

            table_rows = []

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    row_data.append(
                        cell.text.strip()
                    )

                table_rows.append(
                    row_data
                )

            extracted_tables.append(
                {
                    "table_index": table_index,
                    "rows": table_rows
                }
            )

        print("\nEXTRACTED TABLES:")
        print(extracted_tables)

        # =====================================================
        # IMAGE RELATIONSHIP BUILDING
        # =====================================================

        print("\nBUILDING IMAGE MAPPINGS\n")

        rid_to_image = {}
        occurrence_to_image = {}

        with zipfile.ZipFile(
            docx_path,
            "r"
        ) as z:

            media_files = sorted(
                [
                    file_name
                    for file_name in z.namelist()
                    if file_name.startswith(
                        "word/media/"
                    )
                ]
            )

        for index, media_file in enumerate(
            media_files,
            start=1
        ):

            image_name = (
                f"image_{index}.png"
            )

            rel_id = f"rId{index + 8}"

            rid_to_image[
                rel_id
            ] = image_name

        print("\nRID TO IMAGE")

        for rid, image_name in rid_to_image.items():

            print(
                rid,
                "->",
                image_name
            )

        occurrence_counter = 0

        print("\nDRAWING OCCURRENCES\n")

        for para in doc.paragraphs:

            for run in para.runs:

                try:

                    drawings = run._element.xpath(
                        ".//*[local-name()='blip']"
                    )

                    for drawing in drawings:

                        embed = drawing.get(
                            qn("r:embed")
                        )

                        occurrence_counter += 1

                        actual_image = (
                            rid_to_image.get(
                                embed
                            )
                        )

                        occurrence_to_image[
                            occurrence_counter
                        ] = actual_image

                        print(
                            f"OCCURRENCE "
                            f"{occurrence_counter}"
                            f" -> {embed}"
                            f" -> {actual_image}"
                        )

                except Exception:
                    pass

        print(
            "\nTOTAL OCCURRENCES:",
            occurrence_counter
        )

        # =====================================================
        # REMOVE HEADER / FOOTER
        # =====================================================

        for section in doc.sections:

            for p in section.header.paragraphs:
                p.text = ""

            for p in section.footer.paragraphs:
                p.text = ""

        doc.save(docx_path)

        # =====================================================
        # DEBUG DOCX CONTENT
        # =====================================================

        print("\n========================")
        print("DOCX INTERNAL FILES")
        print("========================")

        with zipfile.ZipFile(
            docx_path,
            "r"
        ) as z:

            for name in z.namelist():

                if (
                    "media" in name.lower()
                    or "drawing" in name.lower()
                    or "chart" in name.lower()
                    or "embeddings" in name.lower()
                ):
                    print(name)

        # =====================================================
        # IMAGE EXTRACTION
        # =====================================================

        image_dir = docx_path.replace(
            ".docx",
            "_images"
        )

        os.makedirs(
            image_dir,
            exist_ok=True
        )

        image_paths = []

        with zipfile.ZipFile(
            docx_path,
            "r"
        ) as z:

            media_files = sorted(
                [
                    file_name
                    for file_name in z.namelist()
                    if file_name.startswith(
                        "word/media/"
                    )
                ]
            )

            print("\nMEDIA FILES FOUND:")
            print(len(media_files))

            for index, media_file in enumerate(
                media_files,
                start=1
            ):

                image_data = z.read(
                    media_file
                )

                extension = os.path.splitext(
                    media_file
                )[1]

                image_name = (
                    f"image_{index}{extension}"
                )

                image_path = os.path.join(
                    image_dir,
                    image_name
                )

                with open(
                    image_path,
                    "wb"
                ) as f:

                    f.write(
                        image_data
                    )

                image_paths.append(
                    image_path
                )

        print("\nIMAGES FOUND:")
        print(len(image_paths))

        for img in image_paths:
            print(img)

        # =====================================================
        # TEXT + IMAGE PLACEHOLDERS
        # =====================================================

        text_lines = []

        image_counter = 0

        for para in doc.paragraphs:

            para_text = para.text.strip()

            if para_text:

                text_lines.append(
                    para_text
                )

            xml = para._element.xml

            if "graphicData" in xml:

                image_counter += 1

                actual_image = (
                    occurrence_to_image.get(
                        image_counter
                    )
                )

                print(
                    "PLACEHOLDER CREATED:",
                    image_counter,
                    "->",
                    actual_image
                )

                if actual_image:

                    text_lines.append(
                        f"[[IMAGE:{actual_image}]]"
                    )

        cleaned_text = "\n".join(
            text_lines
        )

        noise_patterns = [
            "adda247.com/defence",
            "www.sscadda.com",
            "www.bankersadda.com",
            "www.adda247.com",
        ]

        for pattern in noise_patterns:

            cleaned_text = cleaned_text.replace(
                pattern,
                ""
            )

        print("\nTEXT SAMPLE:")
        print(
            cleaned_text[:5000]
        )

        # =====================================================
        # QUESTION -> IMAGE MAPPING
        # (DOES NOT TOUCH GROUP LOGIC)
        # =====================================================

        question_images = {}

        pending_images = []

        current_question = None

        for line in text_lines:

            line = line.strip()

            if not line:
                continue

            if line.startswith("[[IMAGE:"):

                image_name = (
                    line.replace(
                        "[[IMAGE:",
                        ""
                    )
                    .replace(
                        "]]",
                        ""
                    )
                    .strip()
                )

                pending_images.append(
                    image_name
                )

                continue

            q_match = re.match(
                r"Q\s*(\d+)\.",
                line,
                re.IGNORECASE
            )

            if q_match:

                question_number = int(
                    q_match.group(1)
                )

                current_question = (
                    question_number
                )

                if (
                    question_number
                    not in question_images
                ):
                    question_images[
                        question_number
                    ] = []

                if pending_images:

                    question_images[
                        question_number
                    ].extend(
                        pending_images
                    )

                    pending_images.clear()

        print("\nQUESTION IMAGE MAP")
        print(question_images)

        # =====================================================
        # QUESTION -> TABLE MAPPING
        # (STORED ONLY, NO GROUP LOGIC CHANGES)
        # =====================================================

        question_tables = {}

        for table in extracted_tables:

            question_tables[
                table["table_index"]
            ] = table

        # =====================================================
        # RETURN EVERYTHING
        # =====================================================

        return {
            "docx": docx_path,
            "text": cleaned_text,
            "images": image_paths,
            "tables": extracted_tables,
            "question_images": question_images,
            "question_tables": question_tables,
            "rid_to_image": rid_to_image,
            "occurrence_to_image": occurrence_to_image
        }