from pdf2docx import Converter
from docx import Document
from pathlib import Path
import zipfile
import os
def extract_tables(docx_path):
    doc = Document(str(docx_path))

    tables = []

    for table_index, table in enumerate(doc.tables):
        rows = []

        for row in table.rows:
            rows.append(
                [cell.text.strip() for cell in row.cells]
            )

        tables.append({
            "table_index": table_index,
            "rows": rows
        })

    return tables

def process_pdf(pdf_path):
    pdf_path = Path(pdf_path)

    docx_path = pdf_path.with_suffix(".docx")

    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()

    doc = Document(str(docx_path))

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:
            paragraph.text = ""

        for paragraph in footer.paragraphs:
            paragraph.text = ""

    doc.save(str(docx_path))

    extracted_text = []

    doc = Document(str(docx_path))

    for para in doc.paragraphs:
        text = para.text.strip()

        if text:
            extracted_text.append(text)
    extracted_tables = extract_tables(docx_path)
    image_dir = str(docx_path).replace(
        ".docx",
        "_images"
    )

    os.makedirs(
        image_dir,
        exist_ok=True
    )

    image_paths = []

    try:
        with zipfile.ZipFile(
            str(docx_path),
            "r"
        ) as docx_zip:

            media_files = [
                f
                for f in docx_zip.namelist()
                if f.startswith("word/media/")
            ]

            for index, media_file in enumerate(
                media_files,
                start=1
            ):

                extension = os.path.splitext(
                    media_file
                )[1]

                output_file = os.path.join(
                    image_dir,
                    f"image_{index}{extension}"
                )

                with open(
                    output_file,
                    "wb"
                ) as f:

                    f.write(
                        docx_zip.read(media_file)
                    )

                image_paths.append(
                    output_file
                )

    except Exception as e:

        print(
            "IMAGE EXTRACTION ERROR:",
            str(e)
        )

    print("IMAGES EXTRACTED:")
    print(image_paths)

    # Make sure extracted_tables is created before return
    # Example:
    # extracted_tables = extract_tables_from_docx(docx_path)
    print("API RETURN TABLE COUNT:")
    print(len(extracted_tables))

    print("API RETURN TABLES:")
    print(extracted_tables)

    return {
        "pdf_file": str(pdf_path),
        "docx_file": str(docx_path),
        "text": "\n".join(extracted_text),
        "images": image_paths,
        "tables": extracted_tables
    }